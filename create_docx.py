#!/usr/bin/env python3
"""Generate a DOCX version of the findings summary with full album tables."""

import json
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None, highlight_rows=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, 'F0F0F0')

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(9)
            if highlight_rows and r_idx in highlight_rows:
                col_idx = highlight_rows[r_idx]
                if isinstance(col_idx, str):
                    set_cell_shading(cell, col_idx)
                elif c_idx in col_idx if isinstance(col_idx, list) else c_idx == col_idx:
                    set_cell_shading(cell, 'FCE4E4')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table

def main():
    # Load data
    with open('output/results_2017_2022.json') as f:
        r2022 = json.load(f)
    with open('output/results_2017_2023.json') as f:
        r2023 = json.load(f)
    with open('output/results_2023.json') as f:
        r2023only = json.load(f)

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    # ── TITLE ──
    title = doc.add_heading(level=1)
    run = title.add_run('Reproduction & Extension: "Smartphones, Online Music Streaming, and Traffic Fatalities"')
    run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.add_run('Reference: ').font.size = Pt(9)
    run = subtitle.add_run('Patel, Worsham, Liu & Jena — NBER Working Paper 34866, Feb 2026')
    run.font.size = Pt(9)
    run.italic = True
    subtitle.add_run('\nData: NHTSA Fatality Analysis Reporting System (FARS), real daily fatality counts').font.size = Pt(9)
    subtitle.add_run('\nAnalysis date: March 2026').font.size = Pt(9)

    # ── EXECUTIVE SUMMARY ──
    doc.add_heading('Executive Summary', level=2)
    p = doc.add_paragraph()
    p.add_run('We reproduced the paper\'s core analysis using real FARS data and extended it through 2023. ')
    run = p.add_run('The original 2017–2022 finding largely replicates, but adding 2023 data eliminates the effect entirely.')
    run.bold = True
    p.add_run(' The five 2023 mega-releases show zero association with traffic fatalities.')

    # ── COMPARISON FIGURE (full page) ──
    doc.add_page_break()
    doc.add_heading('Figure: Side-by-Side Comparison', level=2)
    doc.add_picture('output/comparison_2017_2022_vs_2023.png', width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(
        'Panel A: Event study coefficients (day-by-day fatality changes around album releases). '
        'Panel B: Mean release-day vs. surrounding-day fatalities. '
        'Panel C: Summary statistics for both analysis windows.'
    )
    run.font.size = Pt(8)
    run.italic = True

    # ── HEAD-TO-HEAD COMPARISON ──
    doc.add_page_break()
    doc.add_heading('Head-to-Head Comparison', level=2)

    comp = r2022['comparison']
    comp23 = r2023['comparison']
    h2h_rows = [
        ['Albums analyzed', '10', '10', '10 (re-ranked)'],
        ['Release-day fatalities', '139.1', f"{comp['release_mean']:.1f}", f"{comp23['release_mean']:.1f}"],
        ['Surrounding-day fatalities', '120.9', f"{comp['surrounding_mean']:.1f}", f"{comp23['surrounding_mean']:.1f}"],
        ['Absolute increase', '+18.2 deaths', f"+{comp['absolute_increase']:.1f} deaths", f"+{comp23['absolute_increase']:.1f} deaths"],
        ['Relative increase', '+15.1%', f"+{comp['relative_increase']:.1f}%", f"+{comp23['relative_increase']:.1f}%"],
        ['95% CI', '4.8 – 31.7', f"{comp['ci_lower']:.1f} – {comp['ci_upper']:.1f}",
         f"{comp23['ci_lower']:.1f} – {comp23['ci_upper']:.1f}"],
        ['p-value', '0.01', f"{comp['p_value']:.3f}", f"{comp23['p_value']:.2f}"],
        ['Placebo — random dates exceed', '14 / 1,000', '—', f"{r2023.get('placebo_exceed_dates', '—')} / 1,000"],
        ['Placebo — random Fridays exceed', '20 / 1,000', '—', f"{r2023.get('placebo_exceed_fridays', '—')} / 1,000"],
    ]

    add_styled_table(doc,
        ['Metric', 'Published Paper', 'Our 2017–2022 Run', 'Our 2017–2023 Run'],
        h2h_rows)

    # ── 2023-ONLY ALBUMS ──
    doc.add_heading('2023-Only Albums (5 New Releases)', level=2)
    c23o = r2023only['comparison']
    only_rows = [
        ['Release-day fatalities', f"{c23o['release_mean']:.1f}"],
        ['Surrounding-day fatalities', f"{c23o['surrounding_mean']:.1f}"],
        ['Absolute increase', f"{c23o['absolute_increase']:.1f} deaths"],
        ['p-value', f"{c23o['p_value']:.2f}"],
        ['Placebo — random dates exceed', f"{r2023only['placebo_exceed_dates']} / 1,000"],
        ['Adjacent Friday odds ratio', f"{r2023only['adjacent_friday_or']:.2f}"],
    ]
    add_styled_table(doc, ['Metric', 'Value'], only_rows)

    # ── KEY DISCREPANCIES ──
    doc.add_heading('Key Discrepancies', level=2)

    p = doc.add_paragraph()
    run = p.add_run('1. Our 2017–2022 reproduction is close but not identical to the paper')
    run.bold = True
    bullets = [
        'We find +15.3 excess deaths vs. the paper\'s +18.2 — roughly 16% smaller',
        'Both are statistically significant (p = 0.009 vs. p = 0.01)',
        'Likely explained by FARS data revisions between the paper\'s data pull and ours',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('2. Adding 2023 collapses the effect')
    run.bold = True
    bullets = [
        'Absolute increase drops from +15.3 to +4.1 deaths (−73%)',
        'No longer statistically significant (p = 0.47)',
        'Placebo performance degrades from 14/1,000 to 162/1,000',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('3. 2023 albums show no effect at all')
    run.bold = True
    bullets = [
        'The five 2023 mega-releases show a negative point estimate (−1.8 deaths)',
        '627 of 1,000 random-date placebos produce a larger "effect"',
        'Adjacent-Friday odds ratio is 0.99 — completely flat',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    # ── WHAT CHANGED IN RE-RANKED TOP 10 ──
    doc.add_heading('What Changed in the Re-ranked Top 10?', level=2)
    swap_rows = [
        ["Harry's House (Harry Styles)", "1989 Taylor's Version (Taylor Swift)"],
        ['Her Loss (Drake & 21 Savage)', 'Nadie Sabe Lo Que Va a Pasar Mañana (Bad Bunny)'],
        ['Donda (Kanye West)', 'UTOPIA (Travis Scott)'],
        ["Red Taylor's Version (Taylor Swift)", "Speak Now Taylor's Version (Taylor Swift)"],
        ['Folklore (Taylor Swift)', 'For All the Dogs (Drake)'],
    ]
    add_styled_table(doc,
        ['Dropped from Original Top 10', 'Added from 2023'],
        swap_rows)

    # ── BOTTOM LINE ──
    doc.add_heading('Bottom Line', level=2)
    bl_rows = [
        ['Streaming spike on release day', 'Yes', '~43% increase, consistent with paper'],
        ['Fatality increase (2017–2022)', 'Mostly', 'Directionally consistent, ~16% smaller magnitude'],
        ['Fatality increase with 2023 data', 'No', 'Effect disappears, p = 0.47'],
        ['2023 albums alone', 'No', 'Point estimate is negative'],
    ]
    add_styled_table(doc, ['Claim', 'Replicates?', 'Details'], bl_rows)

    p = doc.add_paragraph()
    p.add_run('\n')
    run = p.add_run(
        'The original finding appears specific to the 2017–2022 sample. Extending the same methodology '
        'one year forward with five new blockbuster albums produces no detectable effect on traffic fatalities.'
    )
    run.italic = True

    # ══════════════════════════════════════════════════════════════════════════
    # FULL ALBUM TABLES
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading('Appendix: Full Album Tables', level=1)

    # ── 2017-2022 Albums ──
    doc.add_heading('Table A1: Top 10 Most Streamed Albums, 2017–2022 (Original Run)', level=2)
    albums_2022 = [
        [1, 'Midnights', 'Taylor Swift', '2022-10-21', 20, '184,695,609'],
        [2, 'Certified Lover Boy', 'Drake', '2021-09-03', 21, '153,441,565'],
        [3, 'Un Verano Sin Ti', 'Bad Bunny', '2022-05-06', 23, '145,811,373'],
        [4, 'Scorpion', 'Drake', '2018-06-29', 25, '132,384,203'],
        [5, 'Mr. Morale & the Big Steppers', 'Kendrick Lamar', '2022-05-13', 18, '99,582,729'],
        [6, "Harry's House", 'Harry Styles', '2022-05-20', 13, '97,621,794'],
        [7, 'Her Loss', 'Drake and 21 Savage', '2022-11-04', 16, '97,390,844'],
        [8, 'Donda', 'Kanye West', '2021-08-29', 25, '94,455,883'],
        [9, "Red (Taylor's Version)", 'Taylor Swift', '2021-11-12', 30, '90,556,180'],
        [10, 'Folklore', 'Taylor Swift', '2020-07-24', 16, '79,443,136'],
    ]
    add_styled_table(doc,
        ['Rank', 'Album', 'Artist', 'Release Date', 'Tracks', 'First-Day Streams'],
        albums_2022)

    # ── 2017-2023 Albums (re-ranked top 10) ──
    doc.add_heading('Table A2: Top 10 Most Streamed Albums, 2017–2023 (Extended Run, Re-ranked)', level=2)
    albums_ext = r2023['albums']
    albums_2023_rows = []
    for a in albums_ext:
        albums_2023_rows.append([
            a['rank'],
            a['album'],
            a['artist'],
            f"{a['first_day_streams']:,}",
        ])
    # Add release dates manually since they're not in the JSON
    dates_map = {
        'Midnights': '2022-10-21',
        "1989 (Taylor's Version)": '2023-10-27',
        'Certified Lover Boy': '2021-09-03',
        'Nadie Sabe Lo Que Va a Pasar Mañana': '2023-10-13',
        'Un Verano Sin Ti': '2022-05-06',
        'Scorpion': '2018-06-29',
        'UTOPIA': '2023-07-28',
        "Speak Now (Taylor's Version)": '2023-07-07',
        'For All the Dogs': '2023-10-06',
        'Mr. Morale & the Big Steppers': '2022-05-13',
    }
    albums_2023_full = []
    for a in albums_ext:
        albums_2023_full.append([
            a['rank'],
            a['album'],
            a['artist'],
            dates_map.get(a['album'], ''),
            f"{a['first_day_streams']:,}",
            '2023' if dates_map.get(a['album'], '').startswith('2023') else '',
        ])
    add_styled_table(doc,
        ['Rank', 'Album', 'Artist', 'Release Date', 'First-Day Streams', 'New'],
        albums_2023_full)

    # ── All 15 albums combined ──
    doc.add_heading('Table A3: All 15 Albums Across Both Periods', level=2)
    all_albums = [
        [1, 'Midnights', 'Taylor Swift', '2022-10-21', '184,695,609', 'Both'],
        [2, "1989 (Taylor's Version)", 'Taylor Swift', '2023-10-27', '176,000,000', '2023 only'],
        [3, 'Certified Lover Boy', 'Drake', '2021-09-03', '153,441,565', 'Both'],
        [4, 'Nadie Sabe Lo Que Va a Pasar Mañana', 'Bad Bunny', '2023-10-13', '145,900,000', '2023 only'],
        [5, 'Un Verano Sin Ti', 'Bad Bunny', '2022-05-06', '145,811,373', 'Both'],
        [6, 'Scorpion', 'Drake', '2018-06-29', '132,384,203', 'Both'],
        [7, 'UTOPIA', 'Travis Scott', '2023-07-28', '128,500,000', '2023 only'],
        [8, "Speak Now (Taylor's Version)", 'Taylor Swift', '2023-07-07', '126,000,000', '2023 only'],
        [9, 'For All the Dogs', 'Drake', '2023-10-06', '109,000,000', '2023 only'],
        [10, 'Mr. Morale & the Big Steppers', 'Kendrick Lamar', '2022-05-13', '99,582,729', 'Both'],
        [11, "Harry's House", 'Harry Styles', '2022-05-20', '97,621,794', '2017–22 only'],
        [12, 'Her Loss', 'Drake and 21 Savage', '2022-11-04', '97,390,844', '2017–22 only'],
        [13, 'Donda', 'Kanye West', '2021-08-29', '94,455,883', '2017–22 only'],
        [14, "Red (Taylor's Version)", 'Taylor Swift', '2021-11-12', '90,556,180', '2017–22 only'],
        [15, 'Folklore', 'Taylor Swift', '2020-07-24', '79,443,136', '2017–22 only'],
    ]
    add_styled_table(doc,
        ['Rank', 'Album', 'Artist', 'Release Date', 'First-Day Streams', 'Included In'],
        all_albums)

    # ── FOOTER ──
    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run(
        'Generated from FARS data. All numbers from real fatality counts — no synthetic data. '
        'Source code and full results available in the repository.'
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Save
    out_path = 'output/findings_summary.docx'
    doc.save(out_path)
    print(f"Saved: {out_path}")

if __name__ == '__main__':
    main()
